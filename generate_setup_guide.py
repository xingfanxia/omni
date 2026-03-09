#!/usr/bin/env python3
"""Generate Microsoft 365 Connector Setup Guide as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_font(run, size=11, bold=False, italic=False, font_name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if not h.runs or run.text != text:
        h.clear()
        run = h.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Calibri"
    return h


def add_para(doc, text="", bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Light gray background via shading
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F2F2F2",
    })
    shading.append(shd)
    return p


def add_bullet(doc, text, bold_prefix="", size=11):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=size, bold=True)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)
        # Light gray header background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "E6E6E6",
        })
        shading.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10)
    doc.add_paragraph()  # spacing after table


def add_note(doc, text, label="Note:"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{label} ")
    set_font(run, size=10, bold=True, italic=True)
    run = p.add_run(text)
    set_font(run, size=10, italic=True)


def add_checkbox(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("☐  " + text)
    set_font(run, size=11)
    return p


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Make all heading styles black
    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.name = "Calibri"

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Microsoft 365 Connector Setup Guide")
    set_font(run, size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Azure Cloud — Docker Compose on a Single VM")
    set_font(run, size=14)
    add_para(doc, "This guide walks a brand-new customer — starting with no Azure subscription — "
             "through every step required to deploy Omni on an Azure VM and connect it to "
             "Microsoft 365 services (Outlook Mail, Outlook Calendar, OneDrive, and SharePoint).",
             size=11, space_after=12)

    doc.add_page_break()

    # ── Table of Contents ──
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "Prerequisites",
        "Create an Azure Account & Subscription",
        "Install the Azure CLI",
        "Provision an Azure VM",
        "Configure Networking (NSG Rules)",
        "Set Up DNS (Optional but Recommended)",
        "Install Docker & Docker Compose on the VM",
        "Register an App in Microsoft Entra ID",
        "Configure Microsoft Graph API Permissions",
        "Create a Client Secret",
        "Grant Admin Consent",
        "Deploy Omni via Docker Compose",
        "Connect Microsoft 365 Sources in Omni",
        "Verify the Integration",
        "Security Hardening Checklist",
        "Troubleshooting",
    ]
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  {item}")
        set_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── 1. Prerequisites ──
    add_heading_styled(doc, "1. Prerequisites", level=1)
    add_table(doc,
        ["Item", "Details"],
        [
            ["A Microsoft account (personal or work)", "Used to create the Azure subscription"],
            ["A valid credit/debit card", "Required for identity verification, even on the free tier"],
            ["A domain name (optional)", "For TLS/HTTPS — you can start with the VM's public IP"],
            ["Microsoft 365 tenant", "The organization whose data you want to index. You need Global Administrator or Application Administrator role to grant admin consent"],
        ])

    # ── 2. Create an Azure Account & Subscription ──
    add_heading_styled(doc, "2. Create an Azure Account & Subscription", level=1)
    add_para(doc, "If you don't have an Azure account yet:")
    add_bullet(doc, "Go to https://azure.microsoft.com/en-us/pricing/purchase-options/azure-account")
    add_bullet(doc, 'Click "Try Azure for free" or "Pay as you go". The free account gives you $200 credit for 30 days and 12 months of popular services at no cost.')
    add_bullet(doc, "Sign in with your Microsoft account (or create one).")
    add_bullet(doc, "Complete the identity verification (phone + credit card).")
    add_bullet(doc, "Once your subscription is active, note your Subscription ID — you'll need it later.")
    add_note(doc, "For production workloads beyond the trial, switch to a Pay-As-You-Go or Enterprise Agreement subscription.", label="Tip:")

    # ── 3. Install the Azure CLI ──
    add_heading_styled(doc, "3. Install the Azure CLI", level=1)
    add_para(doc, "On your local machine (the machine you'll use to manage Azure):")
    add_para(doc, "macOS:", bold=True)
    add_code_block(doc, "brew install azure-cli")
    add_para(doc, "Linux (Ubuntu/Debian):", bold=True)
    add_code_block(doc, "curl -sL https://aka.ms/InstallAzureCLIDeb | sudo bash")
    add_para(doc, "Windows:", bold=True)
    add_code_block(doc, "winget install Microsoft.AzureCLI")
    add_para(doc, "Then log in:")
    add_code_block(doc, "az login")
    add_para(doc, "This opens a browser window for authentication. After login, confirm your subscription:")
    add_code_block(doc, 'az account show --query "{name:name, id:id, tenantId:tenantId}" -o table')
    add_note(doc, "Note the tenantId — you'll need this when configuring the Microsoft connector.")

    # ── 4. Provision an Azure VM ──
    add_heading_styled(doc, "4. Provision an Azure VM", level=1)
    add_para(doc, "Create a Resource Group", bold=True, size=13)
    add_code_block(doc, "az group create \\\n  --name omni-rg \\\n  --location eastus")
    add_para(doc, "Create the VM", bold=True, size=13)
    add_para(doc, "We recommend Standard_D4s_v5 (4 vCPUs, 16 GB RAM) as a minimum for running all Omni services.")
    add_code_block(doc,
        "az vm create \\\n"
        "  --resource-group omni-rg \\\n"
        "  --name omni-vm \\\n"
        "  --image Ubuntu2404 \\\n"
        "  --size Standard_D4s_v5 \\\n"
        "  --admin-username azureuser \\\n"
        "  --generate-ssh-keys \\\n"
        "  --os-disk-size-gb 128 \\\n"
        "  --public-ip-sku Standard")
    add_para(doc, "Save the output — it contains the publicIpAddress you'll use to SSH in and access the Omni UI.")
    add_para(doc, "SSH into the VM", bold=True, size=13)
    add_code_block(doc, "ssh azureuser@<publicIpAddress>")

    # ── 5. Configure Networking ──
    add_heading_styled(doc, "5. Configure Networking (NSG Rules)", level=1)
    add_para(doc, "Azure creates a Network Security Group (NSG) automatically with your VM. By default it allows SSH (port 22). You need to add rules for HTTP/HTTPS access to the Omni UI.")
    add_para(doc, "Open port 443 (HTTPS — recommended for production)", bold=True, size=12)
    add_code_block(doc,
        "az network nsg rule create \\\n"
        "  --resource-group omni-rg \\\n"
        "  --nsg-name omni-vmNSG \\\n"
        "  --name AllowHTTPS \\\n"
        "  --priority 1001 \\\n"
        "  --destination-port-ranges 443 \\\n"
        "  --protocol Tcp \\\n"
        "  --access Allow \\\n"
        "  --direction Inbound")
    add_para(doc, "Open port 80 (HTTP — for Let's Encrypt or dev testing)", bold=True, size=12)
    add_code_block(doc,
        "az network nsg rule create \\\n"
        "  --resource-group omni-rg \\\n"
        "  --nsg-name omni-vmNSG \\\n"
        "  --name AllowHTTP \\\n"
        "  --priority 1002 \\\n"
        "  --destination-port-ranges 80 \\\n"
        "  --protocol Tcp \\\n"
        "  --access Allow \\\n"
        "  --direction Inbound")
    add_para(doc, "Restrict SSH to your IP (strongly recommended)", bold=True, size=12)
    add_code_block(doc,
        "az network nsg rule update \\\n"
        "  --resource-group omni-rg \\\n"
        "  --nsg-name omni-vmNSG \\\n"
        "  --name default-allow-ssh \\\n"
        "  --source-address-prefixes <YOUR_PUBLIC_IP>/32")
    add_note(doc, "Never expose internal service ports (3001-3004, 4001-4009, 5432, 6379) to the internet. Docker Compose's internal network handles all inter-service communication.", label="Important:")

    # ── 6. DNS ──
    add_heading_styled(doc, "6. Set Up DNS (Optional but Recommended)", level=1)
    add_para(doc, "For production with TLS:")
    add_bullet(doc, "In your DNS provider, create an A record pointing your domain (e.g., omni.yourcompany.com) to the VM's public IP.")
    add_bullet(doc, "Omni's built-in Caddy reverse proxy will automatically provision a TLS certificate via Let's Encrypt.")
    add_para(doc, "If you skip this step, you can access Omni via http://<VM_PUBLIC_IP>:3000 for development/testing.")

    # ── 7. Docker ──
    add_heading_styled(doc, "7. Install Docker & Docker Compose on the VM", level=1)
    add_para(doc, "SSH into the VM and run:")
    add_code_block(doc,
        "# Install Docker\n"
        "curl -fsSL https://get.docker.com | sudo sh\n\n"
        "# Add your user to the docker group\n"
        "sudo usermod -aG docker $USER\n\n"
        "# Log out and back in for group change to take effect\n"
        "exit")
    add_para(doc, "SSH back in, then verify:")
    add_code_block(doc, "docker --version\ndocker compose version")
    add_para(doc, "Both commands should succeed. Docker Compose v2 is bundled with modern Docker installations.")

    # ── 8. Register App ──
    add_heading_styled(doc, "8. Register an App in Microsoft Entra ID", level=1)
    add_para(doc, "This creates the identity that Omni uses to access Microsoft 365 data via the Graph API.")
    add_bullet(doc, "Go to the Azure Portal (https://portal.azure.com).")
    add_bullet(doc, 'Navigate to Microsoft Entra ID > App registrations > New registration.')
    add_bullet(doc, "Fill in the registration form:", bold_prefix="")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run("Name: ")
    set_font(run, bold=True)
    run = p.add_run("Omni - Microsoft 365 Connector\n")
    set_font(run)
    run = p.add_run("Supported account types: ")
    set_font(run, bold=True)
    run = p.add_run("Accounts in this organizational directory only (single-tenant)\n")
    set_font(run)
    run = p.add_run("Redirect URI: ")
    set_font(run, bold=True)
    run = p.add_run("Leave blank (not needed for client credentials flow)")
    set_font(run)
    add_bullet(doc, 'Click "Register".')
    add_para(doc, "On the app's Overview page, copy and save:", bold=True)
    add_bullet(doc, " → this is your client_id", bold_prefix="Application (client) ID")
    add_bullet(doc, " → this is your tenant_id", bold_prefix="Directory (tenant) ID")

    # ── 9. API Permissions ──
    add_heading_styled(doc, "9. Configure Microsoft Graph API Permissions", level=1)
    add_para(doc, "The Omni Microsoft connector requires these Application permissions (not Delegated):")
    add_table(doc,
        ["Permission", "Used For"],
        [
            ["User.Read.All", "Enumerating users in the tenant to iterate their data"],
            ["Files.Read.All", "Reading OneDrive files and SharePoint document libraries"],
            ["Mail.Read", "Reading Outlook inbox messages"],
            ["Calendars.Read", "Reading Outlook calendar events"],
            ["Sites.Read.All", "Enumerating SharePoint sites and their document libraries"],
        ])
    add_para(doc, "Steps:", bold=True)
    add_bullet(doc, "In your app registration, go to API permissions.")
    add_bullet(doc, 'Click "Add a permission" > "Microsoft Graph" > "Application permissions".')
    add_bullet(doc, "Search for and add each of the five permissions listed above.")
    add_bullet(doc, 'Click "Add permissions" after selecting all five.')
    add_note(doc, "These are Application permissions, not Delegated. Application permissions allow Omni to access data for all users in the tenant without requiring each user to sign in.")

    # ── 10. Client Secret ──
    add_heading_styled(doc, "10. Create a Client Secret", level=1)
    add_bullet(doc, 'In your app registration, go to "Certificates & secrets".')
    add_bullet(doc, 'Click "Client secrets" > "New client secret".')
    add_bullet(doc, 'Enter a description (e.g., "Omni connector") and choose an expiration (recommended: 24 months).')
    add_bullet(doc, 'Click "Add".')
    add_bullet(doc, "Immediately copy the secret Value (not the Secret ID) — you will not be able to see it again after leaving this page. This is your client_secret.", bold_prefix="")
    add_note(doc, "Set a calendar reminder to rotate the secret before it expires.", label="Important:")

    # ── 11. Admin Consent ──
    add_heading_styled(doc, "11. Grant Admin Consent", level=1)
    add_para(doc, "Application permissions require a tenant administrator to grant consent:")
    add_bullet(doc, "In your app registration, go to API permissions.")
    add_bullet(doc, 'Click "Grant admin consent for [Your Organization]".')
    add_bullet(doc, 'Click "Yes" to confirm.')
    add_bullet(doc, 'All five permissions should now show a green checkmark under the Status column, reading "Granted for [Your Organization]".')
    add_note(doc, 'If you don\'t see the "Grant admin consent" button, you need Global Administrator or Application Administrator role in the Microsoft 365 tenant.', label="Important:")

    # ── 12. Deploy Omni ──
    add_heading_styled(doc, "12. Deploy Omni via Docker Compose", level=1)
    add_para(doc, "Clone the Repository", bold=True, size=13)
    add_code_block(doc, "git clone https://github.com/getomnico/omni.git\ncd omni")
    add_para(doc, "Configure Environment Variables", bold=True, size=13)
    add_code_block(doc, "cp .env.example .env")
    add_para(doc, "Edit the .env file and set these values:")
    add_code_block(doc,
        "# Enable the Microsoft connector\n"
        "ENABLED_CONNECTORS=web,microsoft\n"
        "COMPOSE_PROFILES=${ENABLED_CONNECTORS}\n\n"
        "# Your domain (or VM public IP for testing)\n"
        "APP_URL=https://omni.yourcompany.com\n"
        "OMNI_DOMAIN=omni.yourcompany.com\n"
        "ACME_EMAIL=admin@yourcompany.com\n\n"
        "# Generate a strong encryption key (32+ chars) and salt (16+ chars)\n"
        "ENCRYPTION_KEY=<generate-a-random-string-of-at-least-32-characters>\n"
        "ENCRYPTION_SALT=<generate-a-random-string-of-at-least-16-characters>\n\n"
        "# Change the default database password\n"
        "DATABASE_PASSWORD=<a-strong-random-password>")
    add_para(doc, "Generate secure random values:")
    add_code_block(doc,
        "# Generate ENCRYPTION_KEY (32+ characters)\n"
        "openssl rand -base64 32\n\n"
        "# Generate ENCRYPTION_SALT (16+ characters)\n"
        "openssl rand -base64 16\n\n"
        "# Generate DATABASE_PASSWORD\n"
        "openssl rand -base64 24")
    add_para(doc, "Start Omni", bold=True, size=13)
    add_code_block(doc, "cd docker\ndocker compose up -d")
    add_para(doc, "Check that everything is running:")
    add_code_block(doc, "docker compose ps")
    add_para(doc, "You should see services including omni-microsoft-connector in the list.")

    # ── 13. Connect Sources ──
    add_heading_styled(doc, "13. Connect Microsoft 365 Sources in Omni", level=1)
    add_bullet(doc, "Open the Omni web UI in your browser (https://omni.yourcompany.com or http://<VM_PUBLIC_IP>:3000).")
    add_bullet(doc, "Complete the initial setup wizard (create your admin account).")
    add_bullet(doc, "Navigate to Settings > Connectors (or Sources).")
    add_bullet(doc, "Click Add Source and select the Microsoft 365 service you want to connect:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    for svc, desc in [
        ("OneDrive", "indexes files from all users' OneDrive"),
        ("Outlook Mail", "indexes inbox messages for all users"),
        ("Outlook Calendar", "indexes calendar events for all users"),
        ("SharePoint", "indexes documents from all SharePoint sites"),
    ]:
        run = p.add_run(f"• {svc} — {desc}\n")
        set_font(run, size=11)

    add_bullet(doc, "Enter the credentials from steps 8-10:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    for field in ["Tenant ID: <your-tenant-id>", "Client ID: <your-client-id>", "Client Secret: <your-client-secret>"]:
        run = p.add_run(f"• {field}\n")
        set_font(run, size=11)
    add_bullet(doc, "Click Save / Connect.")
    add_bullet(doc, "Repeat for each Microsoft 365 service you want to index.")
    add_note(doc, "All four Microsoft source types use the same Azure App Registration credentials.")

    # ── 14. Verify ──
    add_heading_styled(doc, "14. Verify the Integration", level=1)
    add_para(doc, "Test the Connection", bold=True, size=12)
    add_para(doc, "After adding a source, Omni will validate the credentials by calling the Microsoft Graph /organization endpoint. If the credentials are correct, the status will show as connected.")
    add_para(doc, "Trigger a Sync", bold=True, size=12)
    add_para(doc, "The connector manager will automatically schedule an initial full sync. You can also trigger it manually from the UI.")
    add_para(doc, "Monitor Sync Progress", bold=True, size=12)
    add_code_block(doc,
        "# View connector logs\n"
        "docker compose logs -f microsoft-connector\n\n"
        "# View connector manager logs\n"
        "docker compose logs -f connector-manager")
    add_para(doc, "What to Expect", bold=True, size=12)
    add_bullet(doc, " Iterates through all users and indexes their Drive files using delta queries.", bold_prefix="OneDrive:")
    add_bullet(doc, " Indexes inbox messages for all users.", bold_prefix="Outlook Mail:")
    add_bullet(doc, " Indexes events within the configured time window (default: 6 months past to 6 months future).", bold_prefix="Outlook Calendar:")
    add_bullet(doc, " Discovers all sites in the tenant and indexes their document libraries.", bold_prefix="SharePoint:")
    add_para(doc, "The first full sync may take significant time depending on data volume. Subsequent incremental syncs use Microsoft Graph delta tokens and are much faster.")

    # ── 15. Security Checklist ──
    add_heading_styled(doc, "15. Security Hardening Checklist", level=1)
    checklist = [
        "Restrict SSH access to your IP or VPN range only",
        "Never expose internal ports (5432/Postgres, 6379/Redis, 3001-3004, 4001-4009) in the NSG",
        "Use HTTPS with a real domain and TLS certificate",
        "Rotate the client secret before it expires — update the credential in Omni's UI when you do",
        "Use strong, unique values for ENCRYPTION_KEY, ENCRYPTION_SALT, and DATABASE_PASSWORD",
        "Keep the VM updated: sudo apt update && sudo apt upgrade -y",
        "Enable Azure Disk Encryption for the OS and data disks",
        "Set up Azure Backup for the VM to protect against data loss",
        "Monitor with Azure Monitor — set alerts for VM CPU, memory, and disk usage",
        "Use Application permissions (not Delegated) — minimum required read-only scopes",
    ]
    for item in checklist:
        add_checkbox(doc, item)

    # ── 16. Troubleshooting ──
    add_heading_styled(doc, "16. Troubleshooting", level=1)

    add_para(doc, '"Insufficient privileges" when granting admin consent', bold=True, size=12)
    add_para(doc, "You need the Global Administrator or Application Administrator role in Microsoft Entra ID. Ask your IT admin to grant consent, or have them assign you the required role.")

    add_para(doc, "401 Unauthorized errors in connector logs", bold=True, size=12)
    add_bullet(doc, "Verify tenant_id, client_id, and client_secret are correct.")
    add_bullet(doc, "Ensure admin consent has been granted (green checkmarks in API permissions).")
    add_bullet(doc, "Check that the client secret hasn't expired.")
    add_bullet(doc, "The connector automatically retries once on 401 by refreshing the token.")

    add_para(doc, "403 Forbidden errors", bold=True, size=12)
    add_bullet(doc, "Usually means a required API permission is missing or admin consent wasn't granted.")
    add_bullet(doc, "Go back to step 9 and verify all five permissions show as \"Granted\".")

    add_para(doc, "429 Too Many Requests", bold=True, size=12)
    add_bullet(doc, "Normal under heavy sync load — the connector automatically respects Retry-After headers.")
    add_bullet(doc, "Microsoft Graph has per-tenant throttling limits; large tenants may need patience during the first full sync.")

    add_para(doc, "Connector container not starting", bold=True, size=12)
    add_code_block(doc,
        "# Check if the microsoft profile is enabled\n"
        "docker compose config --services | grep microsoft\n\n"
        "# If missing, verify your .env has:\n"
        "# ENABLED_CONNECTORS=web,microsoft\n"
        "# COMPOSE_PROFILES=${ENABLED_CONNECTORS}")

    add_para(doc, "No data appearing after sync", bold=True, size=12)
    add_bullet(doc, "Check connector logs: docker compose logs microsoft-connector")
    add_bullet(doc, "Verify the connector-manager is routing correctly: docker compose logs connector-manager")
    add_bullet(doc, "Ensure the indexer is running: docker compose logs indexer")
    add_bullet(doc, "Confirm the app registration permissions match the source type.")

    # ── Quick Reference ──
    doc.add_page_break()
    add_heading_styled(doc, "Quick Reference", level=1)
    add_table(doc,
        ["Item", "Where to Find It"],
        [
            ["Tenant ID", "Azure Portal > Microsoft Entra ID > Overview, or 'az account show'"],
            ["Client ID", "Azure Portal > App registrations > Your app > Overview"],
            ["Client Secret", "Azure Portal > App registrations > Your app > Certificates & secrets"],
            ["VM Public IP", "Azure Portal > Virtual machines > Your VM > Overview"],
            ["Encryption Key", "Generate with: openssl rand -base64 32"],
            ["Encryption Salt", "Generate with: openssl rand -base64 16"],
        ])

    # ── Save ──
    output_path = "/home/user/omni/Microsoft_365_Connector_Setup_Guide.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    build_document()
